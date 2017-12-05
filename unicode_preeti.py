#! /usr/bin/env python3
"""
	Author: Prashant Subedi
	This script is meant to be used as a part of file to file conversion function for unicode to preeti convertor.
"""
#for paragraphs
from docx import Document
raw_file=open("raw_file.nct","r",encoding="utf-8")
a=raw_file.read()
docx_name=a[6:len(a)]
print(docx_name)
raw_file.close()
raw_file=open("raw_file.nct","w",encoding="utf-8")
#Closing the file in read mode and opening it in the write mode
document=Document(docx_name)#This string will be replaced by docx_name
i=0
for paragraph in document.paragraphs:
	for run in paragraph.runs:
			if len(run.text)>0:
				raw_file.write("%s%s"%(run.text,"\0"))
				i+=1
for table in document.tables:
	for row in table.rows: 
		for cell in row.cells:
   			for paragraph in cell.paragraphs:
   				for run in paragraph.runs:
   					if len(run.text)>0:
   						raw_file.write("%s%s"%(run.text,"\0"))
   						i+=1
raw_file.close()
